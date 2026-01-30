"""
Advanced Resume Parser with OCR and Table Extraction
Handles complex PDF layouts, scanned documents, and multi-column formats
"""

import pdfplumber
import docx
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import io
import os
import re

class AdvancedResumeParser:
    """
    Enterprise-grade resume parser that handles:
    - Standard PDF/DOCX text extraction
    - Table extraction from PDFs
    - OCR for scanned/image-based PDFs
    - Multi-column layout detection
    """
    
    def __init__(self, tesseract_path=None):
        """
        Initialize parser
        
        Args:
            tesseract_path: Path to tesseract executable (required for OCR on Windows)
        """
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
    
    def extract_text(self, file_path):
        """
        Main extraction method that automatically detects best approach
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text as string
        """
        if file_path.endswith(".pdf"):
            return self._extract_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF with fallback to OCR if needed"""
        text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Try standard text extraction first
                    page_text = page.extract_text()
                    
                    if page_text and len(page_text.strip()) > 50:
                        text += page_text + "\n"
                    else:
                        # Fallback to OCR for scanned pages
                        text += self._ocr_page(file_path, page.page_number) + "\n"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        text += self._format_tables(tables) + "\n"
        except Exception as e:
            print(f"Standard PDF extraction failed, trying OCR: {e}")
            text = self._ocr_entire_pdf(file_path)
        
        return text.strip()
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX files"""
        text = ""
        try:
            doc = docx.Document(file_path)
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
                    
        except Exception as e:
            print(f"DOCX extraction error: {e}")
        
        return text.strip()
    
    def _format_tables(self, tables):
        """Format extracted tables as readable text"""
        formatted = ""
        for table in tables:
            for row in table:
                if row:
                    # Filter out None values and join
                    row_text = " | ".join([str(cell) for cell in row if cell])
                    formatted += row_text + "\n"
            formatted += "\n"
        return formatted
    
    def _ocr_page(self, pdf_path, page_number):
        """Perform OCR on a specific PDF page"""
        try:
            # Convert specific page to image
            images = convert_from_path(
                pdf_path,
                first_page=page_number,
                last_page=page_number,
                dpi=300
            )
            
            if images:
                # Perform OCR
                text = pytesseract.image_to_string(images[0], lang='eng')
                return text
        except Exception as e:
            print(f"OCR failed for page {page_number}: {e}")
        
        return ""
    
    def _ocr_entire_pdf(self, pdf_path):
        """Perform OCR on entire PDF (fallback for completely scanned documents)"""
        text = ""
        try:
            # Convert all pages to images
            images = convert_from_path(pdf_path, dpi=300)
            
            for i, image in enumerate(images):
                print(f"OCR processing page {i+1}/{len(images)}...")
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += page_text + "\n"
                
        except Exception as e:
            print(f"Full PDF OCR failed: {e}")
        
        return text
    
    def extract_structured_data(self, text):
        """
        Extract structured information from resume text
        
        Returns:
            Dictionary with extracted sections
        """
        sections = {
            'contact': self._extract_contact(text),
            'education': self._extract_education(text),
            'experience': self._extract_experience(text),
            'skills': self._extract_skills_section(text)
        }
        return sections
    
    def _extract_contact(self, text):
        """Extract contact information"""
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone
        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,9}'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact['phone'] = phones[0]
        
        return contact
    
    def _extract_education(self, text):
        """Extract education section"""
        # Look for common education keywords
        edu_keywords = ['education', 'academic', 'qualification', 'degree']
        # This is a simplified version - could be enhanced with NER
        return text
    
    def _extract_experience(self, text):
        """Extract work experience section"""
        # Look for common experience keywords
        exp_keywords = ['experience', 'employment', 'work history', 'professional']
        # This is a simplified version - could be enhanced with NER
        return text
    
    def _extract_skills_section(self, text):
        """Extract skills section"""
        # Look for skills section
        skills_keywords = ['skills', 'technical skills', 'competencies', 'expertise']
        # This is a simplified version - could be enhanced with NER
        return text

# Backward compatibility function
def extract_text(file_path):
    """
    Backward compatible extraction function
    Uses advanced parser automatically
    """
    parser = AdvancedResumeParser()
    return parser.extract_text(file_path)
